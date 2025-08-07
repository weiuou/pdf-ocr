#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文本格式化模块
负责文本的格式化、段落整理和输出格式化
"""

import os
import logging
import re
from typing import List, Dict, Any, Optional
from datetime import datetime
from utils import ensure_directory, safe_filename, FileProcessError
from config import config

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TextFormatter:
    """文本格式化器类"""
    
    def __init__(self):
        """初始化文本格式化器"""
        self.logger = logging.getLogger("pdf_ocr.text_formatter")
        self.preserve_formatting = config.preserve_formatting
        self.output_directory = config.output_directory
        ensure_directory(self.output_directory)
    
    def format_ocr_results(self, ocr_results: List[Dict[str, Any]], 
                          preserve_page_breaks: bool = True) -> str:
        """格式化OCR结果
        
        Args:
            ocr_results: OCR结果列表
            preserve_page_breaks: 是否保留分页符
            
        Returns:
            格式化后的文本
        """
        if not ocr_results:
            return ""
        
        formatted_lines = []
        
        for i, result in enumerate(ocr_results):
            page_number = result.get('page_number', i + 1)
            text = result.get('text', '')
            confidence = result.get('confidence', 0)
            
            # 添加页面标识（如果需要）
            if preserve_page_breaks and len(ocr_results) > 1:
                formatted_lines.append(f"\n--- 第 {page_number} 页 ---\n")
            
            # 处理文本内容
            if text.strip():
                # 格式化文本
                formatted_text = self._format_text_content(text)
                formatted_lines.append(formatted_text)
                
                # 添加置信度信息（如果置信度较低）
                if confidence < config.confidence_threshold:
                    formatted_lines.append(f"\n[注意: 此页识别置信度较低 ({confidence:.1f}%)]\n")
            else:
                formatted_lines.append("[此页无法识别文字内容]\n")
        
        return '\n'.join(formatted_lines)
    
    def _format_text_content(self, text: str) -> str:
        """格式化文本内容
        
        Args:
            text: 原始文本
            
        Returns:
            格式化后的文本
        """
        if not text:
            return ""
        
        # 基本清理
        text = text.strip()
        
        if not self.preserve_formatting:
            # 简单格式化：移除多余空白
            text = re.sub(r'\s+', ' ', text)
            return text
        
        # 保持格式化：智能段落处理
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测是否为标题（简单规则）
            if self._is_title_line(line):
                formatted_lines.append(f"\n{line}\n")
            else:
                formatted_lines.append(line)
        
        # 重新组合段落
        return self._merge_paragraphs(formatted_lines)
    
    def _is_title_line(self, line: str) -> bool:
        """判断是否为标题行
        
        Args:
            line: 文本行
            
        Returns:
            是否为标题
        """
        # 简单的标题检测规则
        title_patterns = [
            r'^第[一二三四五六七八九十\d]+章',  # 章节
            r'^第[一二三四五六七八九十\d]+节',  # 节
            r'^\d+\.',  # 数字编号
            r'^[一二三四五六七八九十]、',  # 中文编号
            r'^\([一二三四五六七八九十\d]+\)',  # 括号编号
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, line):
                return True
        
        # 短行且全大写可能是标题
        if len(line) < 50 and line.isupper():
            return True
        
        return False
    
    def _merge_paragraphs(self, lines: List[str]) -> str:
        """合并段落
        
        Args:
            lines: 文本行列表
            
        Returns:
            合并后的文本
        """
        if not lines:
            return ""
        
        paragraphs = []
        current_paragraph = []
        
        for line in lines:
            if not line.strip():
                continue
            
            # 如果是标题行，结束当前段落
            if line.startswith('\n') and line.endswith('\n'):
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
                paragraphs.append(line.strip())
            else:
                current_paragraph.append(line)
        
        # 添加最后一个段落
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        return '\n\n'.join(paragraphs)
    
    def save_as_txt(self, text: str, output_path: str, 
                   add_metadata: bool = True) -> bool:
        """保存为TXT文件
        
        Args:
            text: 文本内容
            output_path: 输出路径
            add_metadata: 是否添加元数据
            
        Returns:
            保存是否成功
        """
        try:
            content = text
            
            if add_metadata:
                metadata = self._generate_metadata()
                content = f"{metadata}\n\n{'-' * 50}\n\n{text}"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.logger.info(f"文本已保存为: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"保存TXT文件失败: {e}")
            return False
    
    def save_as_docx(self, text: str, output_path: str, 
                    add_metadata: bool = True) -> bool:
        """保存为DOCX文件
        
        Args:
            text: 文本内容
            output_path: 输出路径
            add_metadata: 是否添加元数据
            
        Returns:
            保存是否成功
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx库未安装，无法保存DOCX格式")
            return False
        
        try:
            doc = Document()
            
            # 添加元数据
            if add_metadata:
                metadata = self._generate_metadata()
                doc.add_paragraph(metadata)
                doc.add_paragraph('-' * 50)
            
            # 添加文本内容
            paragraphs = text.split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # 检测标题
                    if self._is_title_line(paragraph_text.strip()):
                        doc.add_heading(paragraph_text.strip(), level=2)
                    else:
                        doc.add_paragraph(paragraph_text.strip())
            
            doc.save(output_path)
            self.logger.info(f"文档已保存为: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"保存DOCX文件失败: {e}")
            return False
    
    def _generate_metadata(self) -> str:
        """生成元数据
        
        Returns:
            元数据字符串
        """
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return f"""PDF OCR 文字提取结果
生成时间: {timestamp}
工具版本: PDF OCR Tool v1.0
识别语言: {config.ocr_language}
识别DPI: {config.ocr_dpi}"""
    
    def save_text(self, text: str, output_path: str, 
                 format_type: str = None) -> bool:
        """保存文本到文件
        
        Args:
            text: 文本内容
            output_path: 输出路径
            format_type: 输出格式 ('txt', 'docx')
            
        Returns:
            保存是否成功
        """
        if format_type is None:
            format_type = config.output_format
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            ensure_directory(output_dir)
        
        # 安全化文件名（只处理文件名部分，保留路径）
        dir_path = os.path.dirname(output_path)
        filename = os.path.basename(output_path)
        safe_filename_only = safe_filename(filename)
        safe_path = os.path.join(dir_path, safe_filename_only) if dir_path else safe_filename_only
        
        if format_type.lower() == 'txt':
            return self.save_as_txt(text, safe_path)
        elif format_type.lower() == 'docx':
            return self.save_as_docx(text, safe_path)
        else:
            self.logger.error(f"不支持的输出格式: {format_type}")
            return False
    
    def create_summary_report(self, ocr_results: List[Dict[str, Any]], 
                            output_path: str, processing_time: float = 0) -> bool:
        """创建处理摘要报告
        
        Args:
            ocr_results: OCR结果列表
            output_path: 输出路径
            processing_time: 处理时间（秒）
            
        Returns:
            报告创建是否成功
        """
        try:
            total_pages = len(ocr_results)
            total_chars = sum(r.get('char_count', 0) for r in ocr_results)
            total_words = sum(r.get('word_count', 0) for r in ocr_results)
            avg_confidence = sum(r.get('confidence', 0) for r in ocr_results) / total_pages if total_pages > 0 else 0
            
            # 统计低置信度页面
            low_confidence_pages = [
                r.get('page_number', i+1) 
                for i, r in enumerate(ocr_results) 
                if r.get('confidence', 0) < config.confidence_threshold
            ]
            
            # 生成报告
            report = f"""PDF OCR 处理摘要报告
{'=' * 50}

处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
处理耗时: {processing_time:.2f} 秒

文档统计:
- 总页数: {total_pages}
- 总字符数: {total_chars:,}
- 总词数: {total_words:,}
- 平均置信度: {avg_confidence:.2f}%

质量评估:
- 置信度阈值: {config.confidence_threshold}%
- 低置信度页面数: {len(low_confidence_pages)}
- 低置信度页面: {', '.join(map(str, low_confidence_pages)) if low_confidence_pages else '无'}

配置信息:
- OCR语言: {config.ocr_language}
- 图片DPI: {config.ocr_dpi}
- 输出格式: {config.output_format}
- 保持格式: {'是' if config.preserve_formatting else '否'}

页面详情:
{'-' * 30}
"""
            
            # 添加每页详情
            for result in ocr_results:
                page_num = result.get('page_number', 0)
                confidence = result.get('confidence', 0)
                char_count = result.get('char_count', 0)
                word_count = result.get('word_count', 0)
                
                status = "正常" if confidence >= config.confidence_threshold else "低置信度"
                if 'error' in result:
                    status = "错误"
                
                report += f"第 {page_num} 页: {char_count} 字符, {word_count} 词, 置信度 {confidence:.1f}% ({status})\n"
            
            # 保存报告
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report)
            
            self.logger.info(f"摘要报告已保存: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"创建摘要报告失败: {e}")
            return False
    
    def export_structured_data(self, ocr_results: List[Dict[str, Any]], 
                             output_path: str, format_type: str = 'json') -> bool:
        """导出结构化数据
        
        Args:
            ocr_results: OCR结果列表
            output_path: 输出路径
            format_type: 导出格式 ('json', 'csv')
            
        Returns:
            导出是否成功
        """
        try:
            if format_type.lower() == 'json':
                import json
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(ocr_results, f, ensure_ascii=False, indent=2)
            
            elif format_type.lower() == 'csv':
                import csv
                with open(output_path, 'w', newline='', encoding='utf-8') as f:
                    if ocr_results:
                        writer = csv.DictWriter(f, fieldnames=ocr_results[0].keys())
                        writer.writeheader()
                        writer.writerows(ocr_results)
            
            else:
                self.logger.error(f"不支持的导出格式: {format_type}")
                return False
            
            self.logger.info(f"结构化数据已导出: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"导出结构化数据失败: {e}")
            return False